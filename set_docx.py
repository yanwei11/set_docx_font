#coding=utf-8
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

from xml.dom.minidom import parse
import xml.dom.minidom

import os
import glob
import getopt
import sys

TABEL_PIC_NOTE_SIZE = 30

class paragragh_style:
    pass
#从xml配置文件中读取文档样式配置
"""
参数：
    elemnet : text，pic，table 元素
    paragragh_style : paragragh_style类对象
返回值：
    无
"""
def get_paragragh_seting(elemnet,paragragh_style):
    line_before = element.getElementsByTagName('line_before')[0]
    paragragh_style.line_before = int(line_before.childNodes[0].data)
    print("line_before: %s" % line_before.childNodes[0].data)

    line_after = element.getElementsByTagName('line_after')[0]
    print("line_after: %s" % line_after.childNodes[0].data)
    paragragh_style.line_after = int(line_after.childNodes[0].data)

    line_spacing = element.getElementsByTagName('line_spacing')[0]
    print("line_spacing: %s" % line_spacing.childNodes[0].data)
    paragragh_style.line_spacing = int(line_spacing.childNodes[0].data)

    alignment = element.getElementsByTagName('alignment')[0]
    print("alignment: %s" % alignment.childNodes[0].data)
    paragragh_style.alignment = int(alignment.childNodes[0].data)

    line_spacing_rule = element.getElementsByTagName('line_spacing_rule')[0]
    print("line_spacing_rule: %s" % line_spacing_rule.childNodes[0].data)
    paragragh_style.line_spacing_rule = int(line_spacing_rule.childNodes[0].data)

    run = element.getElementsByTagName('run')[0]
    #print(type(run))
    cn_font = run.getElementsByTagName('cn_font')[0]
    #print(type(cn_font.childNodes[0].data))
    paragragh_style.cn_font = cn_font.childNodes[0].data
    print("cn_font: %s" % cn_font.childNodes[0].data)
    
    en_font = run.getElementsByTagName('en_font')[0]
    print("en_font: %s" % en_font.childNodes[0].data)
    paragragh_style.en_font = en_font.childNodes[0].data

    size = run.getElementsByTagName('size')[0]
    #print(type(size.childNodes[0].data))
    print("size: %s" % size.childNodes[0].data)
    paragragh_style.size = int(size.childNodes[0].data)

#根据文本的长度和开头的字符，判断是否是图或表的说明行
"""
参数：
    text : 文本字符串
    start_str : 开头字符，是'图' 或'表'
    index_number : 索引编号
返回值：
    True or False
"""
def is_table_note(text):
    return  (len(text) < TABEL_PIC_NOTE_SIZE) and  text.startswith('表')

#根据文本的长度和开头的字符，判断是否是图或表的说明行
"""
参数：
    text : 文本字符串
    start_str : 开头字符，是'图' 或'表'
    index_number : 索引编号
返回值：
    True or False
"""
def is_pic_note(text):
    return  (len(text) < TABEL_PIC_NOTE_SIZE) and text.startswith('图') 

#设置表和图文字的索引号
"""
参数：
    text : 文本字符串
    start_str : 开头字符，是'图' 或'表'
    index_number : 索引编号
返回值：
    添加编号后的文本行
"""
def set_index_number(text,start_str,index_number):
    number = ' 0123456789一二三四五六七八九十'
    text_list = list(text)
    #删除字符串开头的‘图’
    del text_list[0]
    #获得列表中第一个不是数字的位置，包括中文大写数字
    pos = 0
    for s in text_list:
        #print('in:',s)
        if s in number:
            pos += 1
        else:
            break
    #拼接‘表’ + 自动排序编号 + 后续第一个不是数字开始的字符串
    text_out = start_str + str(index_number) +' '+ ''.join(text_list[pos:]) 
    return text_out

def set_paragraph(paragraph,en_font,cn_font,size,line_spacing,line_spacing_rule,line_before,line_after,alignment):
    #不处理空字符，包括图片
    if len(paragraph.runs) == 1:
        if paragraph.runs[0].text  == '':
            if paragraph.runs[0]._element.drawing_lst is not None:
                print('这是个照片-------------,自动居中')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
    if line_spacing_rule > 2 :
        paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    else:
        paragraph.paragraph_format.line_spacing = None
        paragraph.paragraph_format.line_spacing_rule = line_spacing_rule
    
    paragraph.paragraph_format.space_before = Pt(line_before)
    paragraph.paragraph_format.space_after = Pt(line_after)
    paragraph.alignment = alignment

    for run in paragraph.runs:
        #print(run.text)
        
        #普通文本
        # 加粗
        #run.font.bold = True
        # 斜体
        #run.font.italic = True
        # 下划线
        #run.font.underline = True
        # 删除线
        #run.font.strike = True
        #删除rFonts，因为w:eastAsiaTheme或w:asciiTheme 会导致字体设置不对，不清楚原因
        
        if run._element.rPr is not None:
            if run._element.rPr.rFonts is not None:
                p = run._element.rPr.rFonts
                p.getparent().remove(p)
        
        # 字号
        run.font.size = Pt(size)
        # 阴影
        #run.font.shadow = True
        # 字体颜色
        #run.font.color.rgb = RGBColor(0,0,0)
        #只设置英文，且重新设置rFonts
        run.font.name= en_font #只能设置英文名称
        #只设置中文
        run._element.rPr.rFonts.set(qn('w:eastAsia'),cn_font) 

def print_help():
    print('*'*20)
    help_list = [
    "使用方法：yanshi -i 输入文件名 [-d]",
    "-i 输入文件名称，没有路径则与可执行文件在同一路径下，否则写绝对路径",
    "输入文件必须是docx格式",
    "-d 可选，删除文件中的空行（只包含回车）",
    "自动创建输出文件，输出文件名是’输入文件名‘+'_'.docx",
    "如果存在与输出文件相同名字的文件，则删除同名文件",
    ]
    for message in help_list:
        print(message)
    print('*'*20)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


if __name__ == '__main__':
    input_file_name = ''
    delete_empty = False
    #解析命令行参数
    argv = sys.argv[1:]
    try:
        options, args = getopt.getopt(argv, "hdi:", ["help", "i=","delete"])
        print(options,args)
    except getopt.GetoptError:
        print_help()
        sys.exit()
    for option, value in options:
        if option in ("-h", "--help"):
            print_help()
        if option in ("-i", "--ip"):
            print("input file is: {0}".format(value))
            input_file_name = value
        if option in ("-d","--delete"):
            print("删除空行（只有回车）")
            delete_empty = True
    if (input_file_name == '') or  ('.docx' not in input_file_name ) :
        print_help()
        sys.exit()
   
    #输入文件名和输出文件名
    if not os.path.exists(input_file_name):
        print("文件不存在：",input_file_name)
        sys.exit()
    print('输入文件',input_file_name,sep=':')
    out_file_name = input_file_name[0:input_file_name.rfind('.')] + "__" + input_file_name[input_file_name.rfind('.'):]
    if  os.path.exists(out_file_name):
        os.remove(out_file_name)
    #读取XML配置文件
    print("解析XML配置文件")
    p_text = paragragh_style()
    p_pic = paragragh_style()
    p_table = paragragh_style()
    # 使用minidom解析器打开 XML 文档
    DOMTree = xml.dom.minidom.parse("config.xml")
    collection = DOMTree.documentElement
    
    element = collection.getElementsByTagName("text")[0]
    get_paragragh_seting(element,p_text)

    element = collection.getElementsByTagName("pic")[0]
    get_paragragh_seting(element,p_pic)

    element = collection.getElementsByTagName("table")[0]
    get_paragragh_seting(element,p_table)

    #print(p_text.line_before)
    #print(p_pic.cn_font)
    #print(p_table.cn_font)
    #打开输入文件    
    wordfile = Document(input_file_name)
    pic_num = 0
    table_num = 0
    #删除空白行
    if delete_empty:
        for paragraph in wordfile.paragraphs:     
            if len(paragraph.runs) == 0:
                print(paragraph.style.name,paragraph.text,len(paragraph.runs),sep=':')
                print("删除空白行")
                paragraph.clear()  #清除文字，并不删除段落，run也可以,
                delete_paragraph(paragraph)
    

    for paragraph in wordfile.paragraphs:
        style_name = paragraph.style.name
        #print(style_name)
        if style_name.startswith('Heading'):
            #print(style_name,paragraph.text,sep=':')
            continue
        else:
            #print('文本',paragraph.text,sep=':')

            if is_table_note(paragraph.text):
                #print("这是表",paragraph.text,sep=':')
                table_num += 1
                paragraph.text = set_index_number(paragraph.text,paragraph.text[0],table_num)
                set_paragraph(paragraph,p_table.en_font,p_table.cn_font,p_table.size,
                    p_table.line_spacing,p_table.line_spacing_rule,p_table.line_before,p_table.line_after,p_table.alignment)
            elif is_pic_note(paragraph.text):
                #print("这是图",paragraph.text,sep=':')
                pic_num += 1
                paragraph.text = set_index_number(paragraph.text,paragraph.text[0],pic_num)
                set_paragraph(paragraph,p_pic.en_font,p_pic.cn_font,p_pic.size,
                    p_pic.line_spacing,p_pic.line_spacing_rule,p_pic.line_before,p_pic.line_after,p_pic.alignment)
            else:
                set_paragraph(paragraph,p_text.en_font,p_text.cn_font,p_text.size,
                    p_text.line_spacing,p_text.line_spacing_rule,p_text.line_before,p_text.line_after,p_text.alignment)

   
    wordfile.save(out_file_name)